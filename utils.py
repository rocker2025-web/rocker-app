# utils.py
import streamlit as st
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from validate_docbr import CPF, CNPJ
import json
import requests
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import uuid

# --- NOVA FUNÇÃO DE LOGIN COM CONTA DE SERVIÇO ---
def login_gdrive():
    """
    Autentica com o Google Drive usando uma Conta de Serviço.
    Ele tenta primeiro usar os segredos do Streamlit (para produção na nuvem).
    Se falhar, ele usa o arquivo local 'service_account.json' (para desenvolvimento).
    """
    gauth = GoogleAuth()
    scope = ["https://www.googleapis.com/auth/drive"]
    try:
        # Tenta autenticar usando os segredos do Streamlit (para quando estiver online)
        gauth.credentials = GoogleAuth.get_credentials_from_json(st.secrets["gdrive_service_account"], scope)
    except (AttributeError, KeyError):
        # Se falhar (rodando localmente), usa o arquivo JSON da conta de serviço
        gauth.ServiceAuth("service_account.json")
        
    return GoogleDrive(gauth)

# --- FUNÇÃO PARA GERAR O CONTRATO EM WORD ---
def gerar_contrato_docx(dados):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    p_format = style.paragraph_format
    p_format.line_spacing = 1.5
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    section = doc.sections[0]
    header = section.header
    p_header = header.paragraphs[0]
    run_header = p_header.add_run()
    try:
        run_header.add_picture('assets/logo.png', width=Inches(2.0))
    except FileNotFoundError:
        p_header.text = "Rocker Equipamentos"
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run(f"CONTRATO DE {dados['tipo_contrato'].upper()} Nº {dados['numero_contrato']}\n")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)

    def add_justified_paragraph(text=''):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def add_clausula_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return p
    
    add_clausula_heading("DAS PARTES\n")
    p_locadora = add_justified_paragraph()
    p_locadora.add_run("LOCADORA: ").bold = True
    p_locadora.add_run("ROCKER LOCAÇÃO DE EQUIPAMENTOS PARA CONSTRUÇÃO LTDA, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 15.413.157/0001-16, com sede na Rua Carlos Adriano Rodrigues da Silva, Q40 L01, Bairro Potecas, São José/SC, CEP 88.107-493, neste ato representada na forma de seu contrato social.")
    
    p_locataria = add_justified_paragraph()
    p_locataria.add_run("LOCATÁRIA: ").bold = True
    if dados['cliente']['tipo_pessoa'] == "Pessoa Jurídica":
        p_locataria.add_run(f"{dados['cliente']['nome_razao_social']}, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº {dados['cliente']['cpf_cnpj']}, com sede na {dados['cliente']['endereco']}, {dados['cliente']['cidade']} - {dados['cliente']['estado']}, CEP: {dados['cliente']['cep']}, neste ato representada por seu representante legal, {dados['cliente']['representante_legal']['nome']}, portador(a) do CPF sob o nº {dados['cliente']['representante_legal']['cpf']}.")
    else:
        p_locataria.add_run(f"{dados['cliente']['nome_razao_social']}, inscrito(a) no CPF sob o nº {dados['cliente']['cpf_cnpj']}, residente e domiciliado(a) na {dados['cliente']['endereco']}, {dados['cliente']['cidade']} - {dados['cliente']['estado']}, CEP: {dados['cliente']['cep']}.")
    
    add_justified_paragraph("\nAs partes acima qualificadas celebram o presente contrato, que se regerá pelas cláusulas e condições a seguir.")
    
    add_clausula_heading("\nCLÁUSULA PRIMEIRA – DO OBJETO")
    add_justified_paragraph("1.1. O objeto deste contrato é a locação do(s) equipamento(s) descrito(s) na Cláusula Segunda, para ser(em) utilizado(s) exclusivamente no endereço da obra informado abaixo.")

    add_clausula_heading("\nCLÁUSULA SEGUNDA – DOS EQUIPAMENTOS, VALORES E CONDIÇÕES")
    add_justified_paragraph("2.1. Equipamentos e Valores da Locação:")

    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = 'Table Grid'
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Qtde'
    hdr_cells[2].text = 'Equipamento'
    hdr_cells[3].text = 'Vlr. Unit. Mensal (R$)'
    hdr_cells[4].text = 'Vlr. Total Mensal (R$)'

    total_locacao_mensal = 0
    for i, item in enumerate(dados['itens_contrato']):
        row_cells = tabela.add_row().cells
        row_cells[0].text = f"2.1.{i+1}"
        row_cells[1].text = str(item['quantidade'])
        row_cells[2].text = f"{item['produto']} COM {item['plataforma']}"
        row_cells[3].text = f"{item['valor_unitario']:.2f}"
        valor_total_item = item['quantidade'] * item['valor_unitario']
        row_cells[4].text = f"{valor_total_item:.2f}"
        total_locacao_mensal += valor_total_item
        
    add_justified_paragraph("\n2.2. Resumo Financeiro:")
    add_justified_paragraph(f"Valor Total da Locação Mensal: R$ {total_locacao_mensal:.2f}")
    add_justified_paragraph(f"Custo de Entrega (Frete): R$ {dados['valor_entrega']:.2f}")
    add_justified_paragraph(f"Custo de Recolha (Frete): R$ {dados['valor_recolha']:.2f}")

    add_justified_paragraph("\n2.3. Contato e Endereço da Obra:")
    add_justified_paragraph(f"Contato Responsável na Obra: {dados['contato_nome']}")
    add_justified_paragraph(f"Telefone: {dados['contato_telefone']}")
    add_justified_paragraph(f"Endereço da Obra: {dados['endereco_obra']}")

    add_clausula_heading("\nCLÁUSULA DÉCIMA SEGUNDA – DO FORO")
    add_justified_paragraph("12.1. Fica eleito o foro central da comarca de São José para dirimir eventuais litígios oriundos deste contrato, se solução amigável não advir.")

    add_justified_paragraph("\nE, por estarem justas e contratadas, as partes firmam o presente instrumento em 2 (duas) vias de igual teor e forma, na presença das duas testemunhas abaixo.")
    
    assinatura_data = doc.add_paragraph(f"\nSão José, {dados['data_assinatura']}.")
    assinatura_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n_________________________________________")
    doc.add_paragraph("ROCKER LOCAÇÃO DE EQUIPAMENTOS LTDA\n(LOCADORA)")
    doc.add_paragraph("\n\n_________________________________________")
    doc.add_paragraph(f"{dados['cliente']['nome_razao_social'].upper()}\n(LOCATÁRIA)")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNÇÃO PARA GERAR A FATURA EM WORD ---
def gerar_fatura_docx(dados_fatura):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    tabela_cabecalho = doc.add_table(rows=1, cols=2)
    tabela_cabecalho.style = 'Table Grid'
    tabela_cabecalho.autofit = False
    tabela_cabecalho.allow_autofit = False
    tabela_cabecalho.columns[0].width = Inches(4.5)
    tabela_cabecalho.columns[1].width = Inches(2.0)
    
    celula_logo = tabela_cabecalho.cell(0, 0)
    paragrafo_logo = celula_logo.paragraphs[0]
    run_logo = paragrafo_logo.add_run()
    try:
        run_logo.add_picture('assets/logo.png', width=Inches(1.8))
    except FileNotFoundError:
        celula_logo.text = "Rocker Equipamentos"
    
    celula_detalhes = tabela_cabecalho.cell(0, 1)
    p_titulo_fatura = celula_detalhes.paragraphs[0]
    p_titulo_fatura.text = "FATURA DE LOCAÇÃO"
    p_titulo_fatura.runs[0].bold = True
    p_titulo_fatura.runs[0].font.size = Pt(14)
    celula_detalhes.add_paragraph(f"Nº da Fatura: {dados_fatura['NUMERO_FATURA']}")
    celula_detalhes.add_paragraph(f"Data de Emissão: {dados_fatura['DATA_EMISSAO']}")
    for p in celula_detalhes.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    tabela_info = doc.add_table(rows=1, cols=2)
    tabela_info.style = 'Table Grid'
    tabela_info.autofit = False
    tabela_info.allow_autofit = False
    tabela_info.columns[0].width = Inches(3.25)
    tabela_info.columns[1].width = Inches(3.25)
    
    celula_locadora = tabela_info.cell(0, 0)
    celula_locadora.add_paragraph("LOCADORA").runs[0].bold = True
    celula_locadora.add_paragraph("ROCKER LOCAÇÃO DE EQUP. PARA CONST. LTDA EPP\nCNPJ: 15.413.157/0001-16\n...")
    
    celula_destinatario = tabela_info.cell(0, 1)
    celula_destinatario.add_paragraph("DESTINATÁRIO").runs[0].bold = True
    celula_destinatario.add_paragraph(f"Razão Social/Nome: {dados_fatura['NOME_CLIENTE']}")
    celula_destinatario.add_paragraph(f"CNPJ/CPF: {dados_fatura['CNPJ_CLIENTE']}")
    celula_destinatario.add_paragraph(f"Endereço: {dados_fatura['ENDERECO_CLIENTE']}")
    doc.add_paragraph()

    tabela_pagamento = doc.add_table(rows=3, cols=2)
    tabela_pagamento.style = 'Table Grid'
    tabela_pagamento.autofit = False
    tabela_pagamento.allow_autofit = False
    tabela_pagamento.columns[0].width = Inches(2.5)
    tabela_pagamento.columns[1].width = Inches(4.0)

    campos_pagamento = {
        "Nº Contrato:": dados_fatura['NUMERO_CONTRATO'],
        "Forma de Pagamento:": dados_fatura['FORMA_PAGAMENTO'],
        "Data de Vencimento:": dados_fatura['DATA_VENCIMENTO']
    }
    i = 0
    for campo, valor in campos_pagamento.items():
        celula_campo = tabela_pagamento.cell(i, 0)
        celula_valor = tabela_pagamento.cell(i, 1)
        celula_campo.paragraphs[0].add_run(campo).bold = True
        celula_valor.paragraphs[0].text = valor
        i += 1
    doc.add_paragraph()

    tabela_itens = doc.add_table(rows=1, cols=2)
    tabela_itens.style = 'Table Grid'
    tabela_itens.autofit = False
    tabela_itens.allow_autofit = False
    tabela_itens.columns[0].width = Inches(5.0)
    tabela_itens.columns[1].width = Inches(1.5)

    hdr_cells = tabela_itens.rows[0].cells
    hdr_cells[0].text = 'Descrição'
    hdr_cells[1].text = 'Valor (R$)'
    row_cells = tabela_itens.add_row().cells
    row_cells[0].text = dados_fatura['DESCRICAO_SERVICO']
    row_cells[1].text = f"{float(dados_fatura['VALOR_TOTAL'].replace(',', '.')):.2f}"
    doc.add_paragraph()

    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_total.add_run(f"Valor Total: R$ {float(dados_fatura['VALOR_TOTAL'].replace(',', '.')):.2f}").bold = True
    
    doc.add_paragraph(f"\nOBSERVAÇÕES: {dados_fatura['OBSERVACAO']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNÇÕES DE GERENCIAMENTO DE NÚMEROS ---
def get_next_contract_number(drive):
    try:
        config_file = get_database_file(drive, "config.json")
        config_data = read_data(config_file)
        ultimo_numero = config_data.get("ultimo_numero_contrato", 0)
        novo_numero = ultimo_numero + 1
        ano_atual = date.today().year
        config_data["ultimo_numero_contrato"] = novo_numero
        write_data(config_file, config_data)
        return f"{str(novo_numero).zfill(5)}-{ano_atual}"
    except Exception as e:
        st.error(f"Erro ao obter número do contrato: {e}")
        return None

def get_next_fatura_number(drive):
    try:
        config_file = get_database_file(drive, "config.json")
        config_data = read_data(config_file)
        ultimo_numero = config_data.get("ultimo_numero_fatura", 0)
        novo_numero = ultimo_numero + 1
        config_data["ultimo_numero_fatura"] = novo_numero
        write_data(config_file, config_data)
        return f"{str(novo_numero).zfill(7)}"
    except Exception as e:
        st.error(f"Erro ao obter número da fatura: {e}")
        return None

# --- FUNÇÕES DE CONEXÃO COM GOOGLE DRIVE ---
def get_database_file(drive, filename):
    file_list = drive.ListFile({'q': f"title='{filename}' and trashed=false"}).GetList()
    if file_list:
        return file_list[0]
    else:
        file = drive.CreateFile({'title': filename, 'mimeType': 'application/json'})
        file.SetContentString('[]')
        file.Upload()
        return file

def read_data(drive_file):
    content = drive_file.GetContentString()
    if not content:
        return []
    return json.loads(content)

def write_data(drive_file, data):
    drive_file.SetContentString(json.dumps(data, indent=4, ensure_ascii=False))
    drive_file.Upload()

# --- FUNÇÕES DE VALIDAÇÃO ---
def validar_e_formatar_cpf(cpf_str):
    cpf = CPF()
    if cpf.validate(cpf_str):
        return cpf.mask(cpf_str)
    return None

def validar_e_formatar_cnpj(cnpj_str):
    cnpj = CNPJ()
    if cnpj.validate(cnpj_str):
        return cnpj.mask(cnpj_str)
    return None

# --- FUNÇÃO DE CONSULTA DE CEP ---
def consultar_cep(cep):
    cep_limpo = "".join(filter(str.isdigit, cep))
    if len(cep_limpo) != 8:
        return None
    try:
        url = f"https://viacep.com.br/ws/{cep_limpo}/json/"
        response = requests.get(url)
        if response.status_code == 200:
            dados = response.json()
            if dados.get("erro"):
                return None
            return dados
        else:
            return None
    except requests.RequestException:
        return None

# --- COMPONENTE DE RODAPÉ ---
def exibir_rodape():
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; font-size: 14px;">
            <p>© 2025 Gerenciamento de Clientes. Todos os direitos reservados.</p>
            <p>Desenvolvido com a expertise da 
                <a href="https://ascendtechdigital.com.br/" target="_blank">AscendTech</a>.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )